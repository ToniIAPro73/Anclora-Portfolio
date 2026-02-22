import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_shading(cell, color):
    """Establece el color de fondo de una celda."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_horizontal_line(paragraph):
    """Añade una línea horizontal debajo de un párrafo."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_premium_manual(input_file, output_file):
    doc = Document()
    
    # --- CONFIGURACIÓN DE ESTILOS ---
    styles = doc.styles
    
    # Estilo Título Portada
    style_title = styles.add_style('MyTitle', 1) # 1 = Paragraph style
    style_title.font.name = 'Arial'
    style_title.font.size = Pt(42)
    style_title.font.bold = True
    style_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_title.paragraph_format.space_after = Pt(20)

    # Estilo H1
    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Arial'
    style_h1.font.size = Pt(24)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    style_h1.paragraph_format.space_before = Pt(30)
    style_h1.paragraph_format.space_after = Pt(12)

    # Estilo H2
    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Arial'
    style_h2.font.size = Pt(18)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
    style_h2.paragraph_format.space_before = Pt(24)
    style_h2.paragraph_format.space_after = Pt(8)

    # Estilo H3
    style_h3 = styles['Heading 3']
    style_h3.font.name = 'Arial'
    style_h3.font.size = Pt(14)
    style_h3.font.bold = True
    style_h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style_h3.paragraph_format.space_before = Pt(18)
    style_h3.paragraph_format.space_after = Pt(6)

    # Estilo Cuerpo
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.space_after = Pt(8)

    # Estilo Código/Diagramas
    style_code = styles.add_style('CodeStyle', 1)
    style_code.font.name = 'Courier New'
    style_code.font.size = Pt(9)
    style_code.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style_code.paragraph_format.space_before = Pt(6)
    style_code.paragraph_format.space_after = Pt(6)

    # --- LECTURA Y PROCESAMIENTO DEL ARCHIVO MARKDOWN ---
    
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: No se encuentra el archivo '{input_file}'")
        return

    in_code_block = False
    in_table = False
    table_rows = []

    for line in lines:
        line = line.rstrip()

        # Detectar bloques de código (diagramas ASCII)
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
            else:
                in_code_block = True
                # Opcional: añadir una línea en blanco antes del bloque
                doc.add_paragraph() 
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = style_code
            continue

        # Detectar Tablas Markdown
        if '|' in line and not line.startswith('┌'):
            if not in_table:
                in_table = True
                table_rows = []
            
            # Limpiar la línea de la tabla
            clean_line = line.strip()
            if clean_line.startswith('|'):
                clean_line = clean_line[1:]
            if clean_line.endswith('|'):
                clean_line = clean_line[:-1]
            
            cells = [cell.strip() for cell in clean_line.split('|')]
            
            # Ignorar líneas de separación de tablas (|---|---|)
            if not all(set(c) <= set('-: ') for c in cells):
                table_rows.append(cells)
            continue
        else:
            if in_table:
                # Renderizar tabla acumulada
                if table_rows:
                    num_cols = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(table_rows):
                        for j, cell_text in enumerate(row_data):
                            if j < num_cols:
                                cell = table.rows[i].cells[j]
                                cell.text = cell_text
                                # Estilo cabecera
                                if i == 0:
                                    set_cell_shading(cell, '1E3A5F')
                                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.font.size = Pt(10)
                                else:
                                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
                                    run.font.size = Pt(10)
                                    
                    doc.add_paragraph() # Espacio tras la tabla
                in_table = False
                table_rows = []

        # Procesar Encabezados
        if line.startswith('# '):
            text = line[2:]
            if "ANCLORA NEXUS" in text.upper() or "MANUAL DE USUARIO" in text.upper():
                # Es la portada
                doc.add_paragraph("\n\n\n")
                p = doc.add_paragraph(text, style='MyTitle')
                add_horizontal_line(p)
            else:
                doc.add_heading(text, level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            # H4 lo tratamos como párrafo en negrita
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            run.bold = True
            run.font.size = Pt(12)
        elif line.strip() == '---':
            # Separador visual
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.strip() == '':
            # Espacio entre párrafos
            # doc.add_paragraph() # Descomentar si desea más espaciado
            pass
        else:
            # Párrafo normal con negritas inline
            p = doc.add_paragraph()
            # Parser simple para negritas (**texto**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Guardar documento
    doc.save(output_file)
    print(f"✅ Documento generado exitosamente: {output_file}")

# --- EJECUCIÓN ---
# Asegúrese de que 'manual_anclora.md' existe en la misma carpeta
create_premium_manual('manual_anclora.md', 'Manual_Anclora_Nexus_Final.docx')